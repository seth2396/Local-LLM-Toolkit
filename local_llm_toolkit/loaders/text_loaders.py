from pathlib import Path
from typing import Callable
import pdfplumber
from docx import Document as DocxDocument
from bs4 import BeautifulSoup
import re

from .Document import Document
from .BaseLoader import BaseLoader
from ..ingesters import FileItem, WebItem


# ── Cleaning config ────────────────────────────────────────────────────────────
# Three separate dicts keyed by doctype. _clean_for_doctype() applies all three in order:
#   1. _ARTIFACTS   — literal string replacements       {find: replace}
#   2. _CLEANING_PATTERNS — regex replacements          [(pattern, replacement)]
#   3. _TRANSFORMS  — callable transforms               [Callable[[str], str]]

_ARTIFACTS: dict[str, dict[str, str]] = {
    "pdf":     {"\xa0": " ", "\f": "\n"},  # non-breaking space, form feed
    "docx":    {"\xa0": " "},
    "html":    {"\xa0": " "},
    "txt":     {"\xa0": " "},
    "md":      {},
    "default": {"\xa0": " "},
}

_CLEANING_PATTERNS: dict[str, list[tuple[str, str]]] = {
    "pdf": [
        (r"-\n",    ""),        # rejoin hyphenated line breaks
        (r"[ \t]+", " "),       # collapse horizontal whitespace
        (r"\n{3,}", "\n\n"),    # collapse excessive blank lines
    ],
    "docx": [
        (r"[ \t]+", " "),
        (r"\n{3,}", "\n\n"),
    ],
    "html": [
        (r"[ \t]+", " "),
        (r"\n{3,}", "\n\n"),
    ],
    "txt": [
        (r"[ \t]+", " "),
        (r"\n{3,}", "\n\n"),
    ],
    "md":      [],  # preserve markdown structure as-is
    "default": [
        (r"[ \t]+", " "),
        (r"\n{3,}", "\n\n"),
    ],
}

_TRANSFORMS: dict[str, list[Callable[[str], str]]] = {
    "pdf":     [],
    "docx":    [],
    "html":    [],
    "txt":     [],
    "md":      [],
    "default": [],
}


def _clean(
    text: str,
    artifacts: dict[str, str] | None = None,
    patterns: list[tuple[str, str]] | None = None,
    transforms: list[Callable[[str], str]] | None = None,
) -> str:
    """
    Clean a text string in three ordered steps.

    Params:
        text: The raw string to clean.
        artifacts: Literal string replacements applied first. {find: replace}
        patterns: Regex replacements applied second. [(pattern, replacement)]
        transforms: Callables applied last for anything that doesn't fit regex. [fn(str) -> str]

    Returns:
        Cleaned and stripped string.
    """
    if artifacts:
        for symbol, replacement in artifacts.items():
            text = text.replace(symbol, replacement)

    if patterns:
        for pattern, replacement in patterns:
            text = re.sub(pattern, replacement, text)

    if transforms:
        for transform in transforms:
            text = transform(text)

    return text.strip()


def _clean_for_doctype(text: str, doctype: str) -> str:
    """Look up all three cleaning configs for the given doctype and apply them."""
    return _clean(
        text,
        artifacts=_ARTIFACTS.get(doctype, _ARTIFACTS["default"]),
        patterns=_CLEANING_PATTERNS.get(doctype, _CLEANING_PATTERNS["default"]),
        transforms=_TRANSFORMS.get(doctype, _TRANSFORMS["default"]),
    )


class PdfLoader(BaseLoader):
    """Loads PDF files, extracting and normalizing plain text from all pages."""

    def load(self, item: FileItem) -> Document:
        """Extract text from all pages of a PDF, clean, and merge PDF metadata."""
        document = Document(item)
        with pdfplumber.open(item.path) as pdf:
            content = "\n".join(
                text for page in pdf.pages
                if (text := page.extract_text())
            )
            document.content = _clean_for_doctype(content, "pdf")
            if pdf.metadata:
                document.metadata.update({
                    k: v for k, v in pdf.metadata.items()
                    if isinstance(v, (str, int, float, bool))
                })
        return document


class DocxLoader(BaseLoader):
    """Loads .docx files, extracting paragraph text and rendering tables as pipe-delimited rows."""

    def extract_table(self, table):
        """Recursively extract text from a python-docx table."""
        rows_output = []

        for row in table.rows:
            cell_texts = []

            for cell in row.cells:
                parts = []

                # Extract paragraph text
                for para in cell.paragraphs:
                    if para.text.strip():
                        parts.append(para.text.strip())

                # Extract nested tables recursively
                for nested in cell.tables:
                    nested_text = self.extract_table(nested)
                    if nested_text:
                        parts.append(nested_text)

                cell_texts.append("\n".join(parts).strip())

            rows_output.append(" | ".join(cell_texts))

        return "\n".join(rows_output)

    def load(self, item: FileItem) -> Document:
        """Extract paragraphs and tables from a .docx file into a single text block."""
        document = Document(item)

        word_doc = DocxDocument(item.path)
        raw = "\n".join([para.text for para in word_doc.paragraphs])

        for table in word_doc.tables:
            raw += "\n" + self.extract_table(table)

        document.content = _clean_for_doctype(raw, "docx")
        return document


class TextLoader(BaseLoader):
    """Loads plain text files (.txt) as-is without any processing."""

    def load(self, item: FileItem) -> Document:
        """Read the full contents of a plain text file."""
        document = Document(item)
        with open(item.path, 'r', encoding='utf-8') as f:
            document.content = _clean_for_doctype(f.read(), "txt")
        return document


class MarkdownLoader(BaseLoader):
    """Loads Markdown files as raw text, preserving all markup for downstream processing."""

    def load(self, item: FileItem) -> Document:
        """Read the full contents of a Markdown file."""
        document = Document(item)
        with open(item.path, 'r', encoding='utf-8') as f:
            document.content = _clean_for_doctype(f.read(), "md")
        return document


class HTMLLoader(BaseLoader):
    """
    Loads HTML content from a local file or a pre-fetched WebItem.

    If the item has an html_content attribute populated (e.g. from WebIngester),
    that content is used directly to avoid a redundant HTTP request. Otherwise,
    HTML is read from the local file path. BeautifulSoup strips all tags,
    returning plain extracted text.
    """

    def load(self, item: FileItem | WebItem) -> Document:
        """Parse HTML and return plain extracted text."""
        document = Document(item)
        if hasattr(item, 'html_content') and item.html_content is not None:
            html = item.html_content
        else:
            with open(item.path, 'r', encoding='utf-8') as f:
                html = f.read()
        soup = BeautifulSoup(html, 'html.parser')
        document.content = _clean_for_doctype(soup.get_text(), "html")
        return document


_TEXT_LOADERS = {
    '.pdf': PdfLoader,
    '.docx': DocxLoader,
    '.txt': TextLoader,
    '.md': MarkdownLoader,
    '.html': HTMLLoader
}
SUPPORTED_TEXT_FORMATS = _TEXT_LOADERS.keys()


class TextDocLoader:
    f"""
    A loader class to handle loading of various text document formats.
    Currently supported formats: {', '.join(SUPPORTED_TEXT_FORMATS)}
    """
    LOADERS = _TEXT_LOADERS

    def load(self, file) -> Document:
        extension = Path(file.path).suffix.lower()
        if extension in self.LOADERS:
            loader = self.LOADERS[extension]
            return loader.load(file.path)

        raise NotImplementedError(
            f"No document loader implemented for file type '{extension}'. "
            "Please register a loader in BaseLoader.LOADERS."
        )
