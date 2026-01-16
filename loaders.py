from PyPDF2 import PdfReader
from pathlib import Path
from docx import Document as DocxDocument
from bs4 import BeautifulSoup
from ingesters import FileItem
import pandas
import json
import os
import re


class Document:
    def __init__(self, file):
        self.content = ""
        self.metadata = {
                "file_name": file.name,
                "extension": file.ext,
                "doctype": file.doctype,
                "base_path": os.path.dirname(file.path),
                "last_modified": file.modified_ts,
                "file_size": file.size_bytes
            }

        def __str__(self):
            """
            Return a trimmed string representation of the Document.
            """
            trim_len = min(len(self.content), 100)
            return f"Document(doctype={self.doctype}, metadata={self.metadata}, content='{self.content[:trim_len]}...')"

        def to_dict(self, max_content_len: int = 100) -> dict:
            """
                Return a dict representation with trimmed content.
            """
            return {
                "extension": self.doctype,
                "metadata": self.metadata,         
                "content": self.content[:max_content_len],
                "length": len(self.content)
            }

class BaseLoader:
    def load(self, file: FileItem) -> Document:
        raise NotImplementedError("Load method not implemented.")

# Loaders for various text document types.
class PdfLoader(BaseLoader):
    def normalize_text(self, raw_text: str) -> str:
        """
            Normalizes text for embedding usin regular expressions
        
            - Replace non-breaking spaces with normal spaces
            - Replace non-breaking spaces with normal spaces.
            - Remove control characters.
            - Collapse multiple spaces/newlines.
        """
        text = raw_text.replace("\xa0", " ")
        text = re.sub(r"\s+", " ", text)  # collapse whitespace
        return text

    def load(self, file: FileItem) -> Document:
        document = Document(file)
        with open(file.path, 'rb') as openfile:
            reader = PdfReader(openfile)
            content = ""
            for page in reader.pages:
                content += page.extract_text()

            #add pdf content to document
            document.content = self.normalize_text(content)
            #add pdf metadata
            document.metadata.update(reader.metadata)
        return document
    
class DocxLoader(BaseLoader):
    def extract_table(self, table):
        """Recursively extract text from a python-docx table."""
        rows_output = []

        for row in table.rows:
            cell_texts = []

            for cell in row.cells:
                parts = []

                # Extract paragraph text
                for para in cell.paragraphs:
                    if para.text.strip():
                        parts.append(para.text.strip())

                # Extract nested tables recursively
                for nested in cell.tables:
                    nested_text = self.extract_table(nested)
                    if nested_text:
                        parts.append(nested_text)

                # Join everything inside this cell
                cell_texts.append("\n".join(parts).strip())

            # Join cells in the row with a separator
            rows_output.append(" | ".join(cell_texts))

        # Join rows with newlines
        return "\n".join(rows_output)

    def load(self, file: FileItem) -> Document:
        #create document
        document = Document(file)

        #update content with word content
        word_doc = DocxDocument(file.path)
        document.content = "\n".join([para.text for para in word_doc.paragraphs])

        #handleing embedded tables
        for table in word_doc.tables:
            document.content += "\n" + self.extract_table(table)

        return document
    
class TextLoader(BaseLoader):
    def load(self, file: FileItem) -> Document:
        document = Document(file)
        with open(file.path, 'r', encoding='utf-8') as file:
            document.content = file.read()
        return document
    
class MarkdownLoader(BaseLoader):
    def load(self, file: FileItem) -> Document:
        document = Document(file)
        with open(file.path, 'r', encoding='utf-8') as file:
            document.content = file.read()
        return document

class HTMLLoader(BaseLoader):
    def load(self, file: FileItem) -> Document:
        document = Document(file)
        with open(path, 'r', encoding='utf-8') as file:
            soup = BeautifulSoup(file, 'html.parser')
            document.content = soup.get_text()
        return document

_TEXT_LOADERS = {'.pdf': PdfLoader, '.docx': DocxLoader, '.txt': TextLoader, '.md': MarkdownLoader, '.html': HTMLLoader}
SUPPORTED_TEXT_FORMATS = _TEXT_LOADERS.keys()


# Loaders for various data document types

class JSONLoader(BaseLoader):
    
    def load(self, file: FileItem) -> Document:
        document = Document(file)
        with open(file.path, 'r', encoding='utf-8') as file:
            document.content = json.load(file)
        return document
    
class CSVLoader(BaseLoader):
    def load(self, file: FileItem) -> Document:
        document = Docment(file)
        document.content = pd.read_csv(file.path)
        return document

class ExcelLoader(BaseLoader):
    def load(self, file: FileItem) -> Document:
        document = Document(file)
        document.content = pd.read_excel(file.path)
        return document

_DATA_LOADERS = {'.json': JSONLoader, '.csv': CSVLoader, '.xlsx': ExcelLoader}
SUPPORTED_DATA_FORMATS = _DATA_LOADERS.keys()

#Image loaders for various image formats. Currently, these are placeholders and need to be implemented

class JPEGLoader(BaseLoader):  #Not implemented yet
    def load(self, file: FileItem) -> Document:
        raise NotImplementedError("Doc type ending with .jpeg not yet implemented.")
    
class PNGLoader(BaseLoader):   #Not implemented yet 
    def load(self, file: FileItem) -> Document:
        raise NotImplementedError("Doc type ending with .png not yet implemented.")
    
class GIFLoader(BaseLoader):   #Not implemented yet
    def load(self, file: FileItem) -> Document:
        raise NotImplementedError("Doc type ending with .gif not yet implemented.")
    
class BMPLoader(BaseLoader):   #Not implemented yet
    def load(self, file: FileItem) -> Document:
        raise NotImplementedError("Doc type ending with .bmp not yet implemented.")

_IMAGE_LOADERS = {}#{'.jpeg': JPEGLoader, '.png': PNGLoader, '.gif': GIFLoader, '.bmp': BMPLoader}
SUPPORTED_IMAGE_FORMATS = _IMAGE_LOADERS.keys()

#Code loaders for various programming languages. Currently, these are placeholders and need to be implemented

class PythonLoader(BaseLoader): #Not implemented yet
    def load(self, file: FileItem) -> Document:
        raise NotImplementedError("Doc type ending with .py not yet implemented.")

class JavaScriptLoader(BaseLoader): #Not implemented yet
    def load(self, file: FileItem) -> Document:
        raise NotImplementedError("Doc type ending with .js not yet implemented.")

class JavaLoader(BaseLoader): #Not implemented yet
    def load(self, file: FileItem) -> Document:
        raise NotImplementedError("Doc type ending with .java not yet implemented.")

class CppLoader(BaseLoader): #Not implemented yet
    def load(self, file: FileItem) -> Document:
        raise NotImplementedError("Doc type ending with .cpp not yet implemented.")

class CSharpLoader(BaseLoader): #Not implemented yet
    def load(self, file: FileItem) -> Document:
        raise NotImplementedError("Doc type ending with .cs not yet implemented.")

_CODE_LOADERS = {} #'.py': PythonLoader, '.js': JavaScriptLoader, '.java': JavaLoader, '.cpp': CppLoader, '.cs': CSharpLoader}
SUPPORTED_CODE_FORMATS = _CODE_LOADERS.keys()


ALL_SUPPORTED_FORMATS = list(SUPPORTED_TEXT_FORMATS) + list(SUPPORTED_DATA_FORMATS) + list(SUPPORTED_IMAGE_FORMATS) + list(SUPPORTED_CODE_FORMATS)


class TextDocLoader:
    f"""
    A loader class to handle loading of various text document formats.
    Currently supported formats: {', '.join(SUPPORTED_TEXT_FORMATS)}
    """
    LOADERS  =  _TEXT_LOADERS 
    def load(self, file: FileItem) -> Document:
        ext = Path(file.path).suffix.lower()
        if ext in self.LOADERS:
            loader = self.LOADERS[ext]
            return loader.load(file.path)

        raise NotImplementedError(
            f"No document loader implemented for file type '{ext}'. "
            "Please register a loader in BaseLoader.LOADERS."
        )

class DataLoader:
    f"""
    A loader class to handle loading of various data document formats.
    Currently supported formats: {', '.join(SUPPORTED_DATA_FORMATS)}
    """
    LOADERS = _DATA_LOADERS
    def load(self, file: FileItem) -> Document:
        ext = Path(file.path).suffix.lower()
        if ext in self.LOADERS:
            loader = self.LOADERS[ext]
            return loader.load(file.path)

        raise NotImplementedError(
            f"No data loader implemented for file type '{ext}'. "
            "Please register a loader in DataLoader.LOADERS."
        )

class ImageLoader:
    f"""
    A loader class to handle loading of various image document formats.
    Currently supported formats: {', '.join(SUPPORTED_IMAGE_FORMATS)}
    """
    LOADERS = _IMAGE_LOADERS
    def load(self, file: FileItem) -> Document:
        ext = Path(file.path).suffix.lower()
        if ext in self.LOADERS:
            loader = self.LOADERS[ext]
            return loader.load(file.path)

        raise NotImplementedError(
            f"No Image loader implemented for file type '{ext}'. "
            "Please register a loader in ImageLoader.LOADERS."
        )

class CodeLoader:
    f"""
    A loader class to handle loading of various code document formats.
    Currently supported formats: {', '.join(SUPPORTED_CODE_FORMATS)}
    """
    LOADERS = {
            '.py': PythonLoader, 
            '.js': JavaScriptLoader,
            '.java': JavaLoader,
            '.cpp': CppLoader,
            '.cs': CSharpLoader
        }
    def load(self, file: FileItem) -> Document:
        ext = Path(file.path).suffix.lower()
        if ext in self.LOADERS:
            loader = self.LOADERS[ext]
            return loader.load(file.path)

        raise NotImplementedError(
            f"No code loader implemented for file type '{ext}'. "
            "Please register a loader in CodeLoader.LOADERS."
        )

class UniversalLoader:
    f"""
    A unified loader class to handle loading of various document formats.
    Currently supported formats: {', '.join(ALL_SUPPORTED_FORMATS)}
    """
    LOADERS = _TEXT_LOADERS | _DATA_LOADERS | _IMAGE_LOADERS | _CODE_LOADERS
    def load(self, file: FileItem) -> Document:
        ext = file.ext
        if ext in self.LOADERS:
            loader_class = self.LOADERS[ext]
            loader = loader_class()
            return loader.load(file)

        raise NotImplementedError(
            f"No loader implemented for file type '{ext}'. "
            "Please register a loader in Loader.loaders."
        )
       
if __name__ == "__main__":
    pass