from pathlib import Path
from PyPDF2 import PdfReader
from docx import Document as DocxDocument
from bs4 import BeautifulSoup
import re

from .Document import Document
from .BaseLoader import BaseLoader


class PdfLoader(BaseLoader):
    def normalize_text(self, raw_text: str) -> str:
        """
            Normalizes text for embedding using regular expressions

            - Replace non-breaking spaces with normal spaces
            - Remove control characters.
            - Collapse multiple spaces/newlines.
        """
        text = raw_text.replace("\xa0", " ")
        text = re.sub(r"\s+", " ", text)  # collapse whitespace
        return text

    def load(self, file) -> Document:
        document = Document(file)
        with open(file.path, 'rb') as openfile:
            reader = PdfReader(openfile)
            content = ""
            for page in reader.pages:
                content += page.extract_text()

            document.content = self.normalize_text(content)
            document.metadata.update(reader.metadata)
        return document


class DocxLoader(BaseLoader):
    def extract_table(self, table):
        """Recursively extract text from a python-docx table."""
        rows_output = []

        for row in table.rows:
            cell_texts = []

            for cell in row.cells:
                parts = []

                # Extract paragraph text
                for para in cell.paragraphs:
                    if para.text.strip():
                        parts.append(para.text.strip())

                # Extract nested tables recursively
                for nested in cell.tables:
                    nested_text = self.extract_table(nested)
                    if nested_text:
                        parts.append(nested_text)

                cell_texts.append("\n".join(parts).strip())

            rows_output.append(" | ".join(cell_texts))

        return "\n".join(rows_output)

    def load(self, file) -> Document:
        document = Document(file)

        word_doc = DocxDocument(file.path)
        document.content = "\n".join([para.text for para in word_doc.paragraphs])

        for table in word_doc.tables:
            document.content += "\n" + self.extract_table(table)

        return document


class TextLoader(BaseLoader):
    def load(self, file) -> Document:
        document = Document(file)
        with open(file.path, 'r', encoding='utf-8') as f:
            document.content = f.read()
        return document


class MarkdownLoader(BaseLoader):
    def load(self, file) -> Document:
        document = Document(file)
        with open(file.path, 'r', encoding='utf-8') as f:
            document.content = f.read()
        return document


class HTMLLoader(BaseLoader):
    def load(self, file) -> Document:
        document = Document(file)
        with open(file.path, 'r', encoding='utf-8') as f:
            soup = BeautifulSoup(f, 'html.parser')
            document.content = soup.get_text()
        return document


_TEXT_LOADERS = {
    '.pdf': PdfLoader,
    '.docx': DocxLoader,
    '.txt': TextLoader,
    '.md': MarkdownLoader,
    '.html': HTMLLoader
}
SUPPORTED_TEXT_FORMATS = _TEXT_LOADERS.keys()


class TextDocLoader:
    f"""
    A loader class to handle loading of various text document formats.
    Currently supported formats: {', '.join(SUPPORTED_TEXT_FORMATS)}
    """
    LOADERS = _TEXT_LOADERS

    def load(self, file) -> Document:
        ext = Path(file.path).suffix.lower()
        if ext in self.LOADERS:
            loader = self.LOADERS[ext]
            return loader.load(file.path)

        raise NotImplementedError(
            f"No document loader implemented for file type '{ext}'. "
            "Please register a loader in BaseLoader.LOADERS."
        )
